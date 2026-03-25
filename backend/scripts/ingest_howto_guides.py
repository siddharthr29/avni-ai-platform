#!/usr/bin/env python3
"""Ingest Avni How-To Guides (.docx, .pdf) into pgvector RAG.

Extracts text from docx/pdf files, chunks them, and indexes into
the 'howto_guides' collection in pgvector.

Usage:
    python scripts/ingest_howto_guides.py [/path/to/guides/dir]

Default path: ~/Downloads/How to guides/
"""

import asyncio
import logging
import os
import re
import sys

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_BACKEND_DIR = os.path.dirname(_SCRIPT_DIR)
sys.path.insert(0, _BACKEND_DIR)

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("ingest_howto")

DATABASE_URL = "postgresql://samanvay@localhost:5432/avni_ai"
COLLECTION = "howto_guides"
DEFAULT_DIR = os.path.expanduser("~/Downloads/How to guides/")


def chunk_text(text: str, max_chars: int = 1500, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks at paragraph boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    paragraphs = text.split("\n\n")
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 > max_chars and current:
            chunks.append(current.strip())
            lines = current.split("\n")
            overlap_text = "\n".join(lines[-3:]) if len(lines) > 3 else current[-overlap:]
            current = overlap_text + "\n\n" + para
        else:
            current = current + "\n\n" + para if current else para
    if current.strip():
        chunks.append(current.strip())
    return chunks


def extract_docx(filepath: str) -> str:
    """Extract text from a .docx file, preserving structure."""
    import docx

    doc = docx.Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        # Preserve heading structure
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            try:
                hashes = "#" * int(level)
            except ValueError:
                hashes = "##"
            parts.append(f"{hashes} {text}")
        else:
            parts.append(text)

    # Also extract table content
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    return "\n\n".join(parts)


def extract_pdf(filepath: str) -> str:
    """Extract text from a .pdf file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        logger.warning("PyPDF2 not installed, skipping PDF: %s", filepath)
        return ""


def clean_title(filename: str) -> str:
    """Extract a clean title from filename."""
    name = os.path.splitext(filename)[0]
    # Remove common prefixes/suffixes
    name = re.sub(r"^How [Tt]o [Gg]uide[_\s-]*", "", name)
    name = re.sub(r"\s*\(\d+\)\s*$", "", name)  # Remove (1) suffix
    name = name.strip("_- ")
    return name or filename


def load_guides(guides_dir: str) -> list[dict]:
    """Load all .docx and .pdf files from the guides directory."""
    chunks = []
    files_processed = 0

    for fname in sorted(os.listdir(guides_dir)):
        filepath = os.path.join(guides_dir, fname)
        if not os.path.isfile(filepath):
            continue

        ext = os.path.splitext(fname)[1].lower()
        if ext == ".docx":
            text = extract_docx(filepath)
        elif ext == ".pdf":
            text = extract_pdf(filepath)
        else:
            continue

        if len(text) < 50:
            logger.warning("Skipping %s (too short: %d chars)", fname, len(text))
            continue

        title = clean_title(fname)
        files_processed += 1

        # Prefix each chunk with the guide title for context
        for i, chunk in enumerate(chunk_text(text, 1500)):
            context_prefix = f"How-To Guide: {title}"
            chunks.append({
                "content": chunk,
                "collection": COLLECTION,
                "context_prefix": context_prefix,
                "metadata": {
                    "title": title,
                    "filename": fname,
                    "chunk": i,
                    "type": "howto_guide",
                },
                "source_file": fname,
            })

    logger.info(
        "How-To Guides: %d chunks from %d files in %s",
        len(chunks), files_processed, guides_dir,
    )
    return chunks


async def ingest(all_chunks: list[dict]):
    """Embed and insert all chunks into pgvector."""
    from app.services.rag.embeddings import EmbeddingClient
    from app.services.rag.vector_store import VectorStore

    emb = EmbeddingClient()
    vs = VectorStore(dsn=DATABASE_URL)
    await vs.initialize()

    logger.info("Clearing existing '%s' collection...", COLLECTION)
    await vs.clear_collection(COLLECTION)

    total = 0
    batch_size = 64
    for i in range(0, len(all_chunks), batch_size):
        batch = all_chunks[i : i + batch_size]
        texts = [c["content"] for c in batch]
        embeddings = emb.embed_batch(texts)
        db_chunks = [
            {
                "collection": COLLECTION,
                "content": c["content"],
                "context_prefix": c.get("context_prefix", ""),
                "embedding": e,
                "metadata": c.get("metadata", {}),
                "source_file": c.get("source_file", ""),
            }
            for c, e in zip(batch, embeddings)
        ]
        await vs.upsert_chunks(db_chunks)
        total += len(db_chunks)
        logger.info("  Ingested %d / %d chunks", total, len(all_chunks))

    # Show final stats
    stats = await vs.get_collection_stats()
    logger.info("=" * 60)
    logger.info("DONE: %d chunks ingested into '%s'", total, COLLECTION)
    logger.info("Total RAG chunks across all collections:")
    grand = 0
    for c, n in sorted(stats.items()):
        logger.info("  %s: %d", c, n)
        grand += n
    logger.info("  TOTAL: %d", grand)
    logger.info("=" * 60)
    await vs.close()


def main():
    guides_dir = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_DIR
    if not os.path.isdir(guides_dir):
        logger.error("Directory not found: %s", guides_dir)
        sys.exit(1)

    all_chunks = load_guides(guides_dir)
    if not all_chunks:
        logger.error("No chunks extracted. Check that .docx/.pdf files exist in %s", guides_dir)
        sys.exit(1)

    asyncio.run(ingest(all_chunks))


if __name__ == "__main__":
    main()
